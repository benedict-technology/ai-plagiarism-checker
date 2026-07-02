"""word_processor.py
Paragraph-by-paragraph plagiarism and AI scan of an uploaded .docx file.
Returns a new annotated .docx with:
  - Red shading    = HIGH risk    (plagiarism >= 70% OR ai >= 75%)
  - Yellow shading = MEDIUM risk  (plagiarism >= 35% OR ai >= 45%)
  - Green shading  = CLEAN
  - A native Word comment on every flagged paragraph listing the top
    source URL as proof, using python-docx's built-in add_comment API.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from io import BytesIO

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from sqlalchemy.orm import Session

from .ai_detection import analyze_ai_content
from .source_tracing import seed_default_sources, trace_sources

_HIGH_COLOUR    = "FF9999"
_MEDIUM_COLOUR  = "FFFF99"
_CLEAN_COLOUR   = "CCFFCC"
_SUMMARY_COLOUR = "D1E8FF"


@dataclass
class ParagraphResult:
    index: int
    text: str
    plagiarism_score: int
    ai_score: int
    ai_label: str
    risk: str
    top_source_url: str | None
    top_source_title: str | None
    top_source_similarity: int


@dataclass
class WordScanResult:
    paragraphs: list[ParagraphResult] = field(default_factory=list)
    annotated_docx_bytes: bytes = b""
    high_count: int = 0
    medium_count: int = 0
    clean_count: int = 0
    error: str = ""


def _risk(plagiarism: int, ai: int) -> str:
    if plagiarism >= 70 or ai >= 75:
        return "high"
    if plagiarism >= 35 or ai >= 45:
        return "medium"
    return "low"


def _hex_colour(risk: str) -> str:
    return {"high": _HIGH_COLOUR, "medium": _MEDIUM_COLOUR}.get(risk, _CLEAN_COLOUR)


def _set_paragraph_shading(paragraph, hex_colour: str) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    for existing in pPr.findall(qn("w:shd")):
        pPr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_colour)
    pPr.append(shd)


def scan_word_document(docx_bytes: bytes, db: Session) -> WordScanResult:
    """Scan every paragraph of a .docx and return an annotated copy."""
    seed_default_sources(db)

    try:
        doc = Document(BytesIO(docx_bytes))
    except Exception as exc:
        return WordScanResult(error=f"Could not open document: {exc}")

    scannable = [
        (i, para)
        for i, para in enumerate(doc.paragraphs)
        if len(para.text.split()) >= 8 and para.runs
    ]

    if not scannable:
        return WordScanResult(
            error=(
                "No scannable paragraphs found. The document may be empty, "
                "contain only headings/short lines, or use unsupported formatting."
            )
        )

    para_results: list[ParagraphResult] = []

    for idx, para in scannable:
        text = para.text.strip()

        try:
            plag_score, _, source_traces = trace_sources(text, db)
        except Exception:
            plag_score, source_traces = 0, []

        try:
            ai_result = analyze_ai_content(text)
            ai_score  = ai_result.score
            ai_label  = ai_result.label
        except Exception:
            ai_score, ai_label = 0, "Unknown"

        risk = _risk(plag_score, ai_score)

        top_url, top_title, top_sim = None, None, 0
        if source_traces:
            top = source_traces[0]
            top_url   = top.url
            top_title = top.title
            top_sim   = top.similarity_percentage

        para_results.append(ParagraphResult(
            index=idx,
            text=text,
            plagiarism_score=plag_score,
            ai_score=ai_score,
            ai_label=ai_label,
            risk=risk,
            top_source_url=top_url,
            top_source_title=top_title,
            top_source_similarity=top_sim,
        ))

        _set_paragraph_shading(para, _hex_colour(risk))

        if risk in {"high", "medium"}:
            lines = [f"Plagiarism: {plag_score}%  |  AI: {ai_score}% ({ai_label})"]
            if top_title:
                lines.append(f"Source: {top_title}")
            if top_url:
                lines.append(f"URL: {top_url}")
            if top_sim:
                lines.append(f"Similarity: {top_sim}%")
            try:
                doc.add_comment(
                    para.runs,
                    text="\n".join(lines),
                    author="AI Plagiarism Checker",
                    initials="APC",
                )
            except Exception:
                pass  # Shading still shows risk even if comment fails

    high_count   = sum(1 for r in para_results if r.risk == "high")
    medium_count = sum(1 for r in para_results if r.risk == "medium")
    clean_count  = sum(1 for r in para_results if r.risk == "low")

    # Summary banner at the top of the document
    summary_text = (
        f"AI PLAGIARISM SCAN SUMMARY  |  "
        f"High risk: {high_count}  |  "
        f"Medium risk: {medium_count}  |  "
        f"Clean: {clean_count}  |  "
        f"Total scanned: {len(para_results)}"
    )
    summary_para = doc.add_paragraph()
    run = summary_para.add_run(summary_text)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    _set_paragraph_shading(summary_para, _SUMMARY_COLOUR)
    doc.paragraphs[0]._p.addprevious(summary_para._p)

    out = BytesIO()
    doc.save(out)

    return WordScanResult(
        paragraphs=para_results,
        annotated_docx_bytes=out.getvalue(),
        high_count=high_count,
        medium_count=medium_count,
        clean_count=clean_count,
    )